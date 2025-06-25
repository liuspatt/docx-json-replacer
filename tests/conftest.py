import pytest
import json
import tempfile
import os
from pathlib import Path
from docx import Document
from docx_json_replacer.docx_replacer import DocxReplacer


@pytest.fixture
def test_data():
    """Fixture providing test data from fixtures/test_data.json"""
    return {
        "input.name": "John Doe",
        "input.date": "2025-06-25",
        "input.company": "Example Corp",
        "input.position": "Software Engineer"
    }


@pytest.fixture
def sample_docx():
    """Fixture providing a temporary DOCX file with placeholders"""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_paragraph("Hello {{input.name}}!")
        doc.add_paragraph("Date: {{input.date}}")
        doc.add_paragraph("Company: {{input.company}}")
        doc.add_paragraph("Position: {{input.position}}")
        
        # Add a table with placeholders
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "{{input.name}}"
        table.cell(1, 0).text = "Position"
        table.cell(1, 1).text = "{{input.position}}"
        
        doc.save(tmp.name)
        yield tmp.name
        os.unlink(tmp.name)


@pytest.fixture
def sample_json_file(test_data):
    """Fixture providing a temporary JSON file with test data"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as tmp:
        json.dump(test_data, tmp, indent=2)
        tmp.flush()
        yield tmp.name
        os.unlink(tmp.name)


@pytest.fixture
def docx_replacer(sample_docx):
    """Fixture providing a DocxReplacer instance with sample document"""
    return DocxReplacer(sample_docx)


@pytest.fixture
def output_docx():
    """Fixture providing a temporary output file path"""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp.close()
        yield tmp.name
        if os.path.exists(tmp.name):
            os.unlink(tmp.name)